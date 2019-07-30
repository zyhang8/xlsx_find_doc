# -*- coding: utf-8 -*-
# @Time    : 2019-06-06 13:22
# @Author  : zyh
# @File    : main.py
# @Software: PyCharm

from openpyxl import load_workbook
from docx import Document

# 由于时间比较急，代码没有那么pythonic，比较粗略，看到不懂再问我吧
def main():
    lists_title = []
    list_docx = [[0 for col in range(1000)] for row in range(1000)]
    list_result = [[0 for col in range(1000)] for row in range(1000)]
    # print(wb.sheetnames)
    wb = load_workbook("session.xlsx")
    sheet = wb["Sheet1"]
    # 创建文档对象
    document_write = Document()

    for i in sheet["A"]:
        # print(i.value, end=" ")
        lists_title.append(str(i.value))
    print(lists_title)
    for i in range(1, len(lists_title), 2):
        print("序号：%s   值：%s" % (i + 1, lists_title[i]))
        # for i in range(1, 864):
        #     path = "All_abstracts_" + str(i) + ".docx"
        #     document = Document(path)
        #     j = 0

    for i in range(0, 834):
        path = "./xlsx_find_docx/All_abstracts_" + str(i + 1) + ".docx"
        # print(path)
        document = Document(path)
        j = 0
        for paragraph in document.paragraphs:
            # print(paragraph.text)
            list_docx[i][j] = paragraph.text
            # print(list_docx[i])
            j += 1
        # print(j)

    for i in range(0, len(lists_title)):
        # xlsx标题的一个
        # 对比docx
        for j in range(0, 834):
            # 第一个docx
            # 遍历行
            for k in range(0, len(list_docx[j])):
                # 如果有相同的
                if lists_title[i] == list_docx[j][k]:
                    list_result[i] = list_docx[j]
                    print(list_result[i])

    for i in range(0, len(lists_title)):
        if list_result[i][0] != 0:
            # document_write.add_paragraph(list_result[i][0])
            document_write.add_paragraph().add_run(list_result[i][0]).bold = True
            print(list_result[i][0])
            for j in range(1, len(list_result[i])):
                if list_result[i][j] != 0:
                    # 往文档中添加段落
                    document_write.add_paragraph(list_result[i][j])
                    print(list_result[i][j])
            # 添加分页符
            document_write.add_page_break()
            # 保存文档
            document_write.save('results.docx')
    # 保存文档
    document_write.save('results.docx')


if __name__ == '__main__':
    main()
